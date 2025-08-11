from __future__ import annotations
import os
from typing import List, Tuple

import chromadb
from chromadb.utils import embedding_functions

from src.config import AppConfig
from src.rag.embeddings import EmbeddingsModel


class RAGIndexer:
    def __init__(self, config: AppConfig):
        self.config = config
        self.client = chromadb.PersistentClient(path=config.vectorstore_dir)
        self.collection_name = "adgm_reference"
        self.emb_model = EmbeddingsModel(config)

    def _get_collection(self):
        # Create or get
        return self.client.get_or_create_collection(
            name=self.collection_name,
            metadata={"hnsw:space": "cosine"},
        )

    def build_or_rebuild(self, force_rebuild: bool = False) -> str:
        collection = self._get_collection()
        if force_rebuild:
            try:
                self.client.delete_collection(self.collection_name)
            except Exception:
                pass
            collection = self._get_collection()

        # If already has items, skip
        count = collection.count()
        if count > 0 and not force_rebuild:
            return f"Index ready (existing {count} items)."

        # Load reference files and URLs
        texts: List[str] = []
        sources: List[str] = []
        for root, _, files in os.walk(self.config.data_reference_dir):
            for fname in files:
                path = os.path.join(root, fname)
                low = fname.lower()
                if low.endswith((".txt", ".md")):
                    with open(path, "r", encoding="utf-8", errors="ignore") as f:
                        content = f.read()
                    for i in range(0, len(content), 1200):
                        chunk = content[i : i + 1200]
                        texts.append(chunk)
                        sources.append(path)
                elif low.endswith(".pdf"):
                    try:
                        from pdfminer.high_level import extract_text
                        content = extract_text(path)
                        for i in range(0, len(content), 1200):
                            chunk = content[i : i + 1200]
                            texts.append(chunk)
                            sources.append(path)
                    except Exception:
                        pass
                elif low.endswith((".docx",)):
                    try:
                        from docx import Document
                        from itertools import chain
                        d = Document(path)
                        paras = [p.text for p in d.paragraphs]
                        for table in d.tables:
                            for row in table.rows:
                                for cell in row.cells:
                                    paras.append(cell.text)
                        content = "\n".join([t for t in paras if t and t.strip()])
                        for i in range(0, len(content), 1200):
                            chunk = content[i : i + 1200]
                            texts.append(chunk)
                            sources.append(path)
                    except Exception:
                        pass

        # Also fetch URLs listed in sources_urls.txt if present
        urls_file = os.path.join(self.config.data_reference_dir, "sources_urls.txt")
        if os.path.exists(urls_file):
            try:
                import requests
                from bs4 import BeautifulSoup
                with open(urls_file, "r", encoding="utf-8", errors="ignore") as f:
                    urls = [u.strip() for u in f.read().splitlines() if u.strip() and not u.strip().startswith("#")]
                for url in urls:
                    try:
                        r = requests.get(url, timeout=20)
                        r.raise_for_status()
                        if url.lower().endswith((".pdf",)):
                            # best-effort: skip binary fetch of large PDFs here (user can pre-download); or store URL as source only
                            texts.append(f"Referenced PDF at {url}")
                            sources.append(url)
                        elif url.lower().endswith((".docx",)):
                            texts.append(f"Referenced DOCX template at {url}")
                            sources.append(url)
                        else:
                            soup = BeautifulSoup(r.text, "html.parser")
                            text = soup.get_text(" ")
                            for i in range(0, len(text), 1200):
                                chunk = text[i : i + 1200]
                                texts.append(chunk)
                                sources.append(url)
                    except Exception:
                        continue
            except Exception:
                pass

        if not texts:
            return "No reference files found. Add files to data/reference/."

        embeddings = self.emb_model.embed(texts)
        ids = [f"ref_{i}" for i in range(len(texts))]
        collection.upsert(ids=ids, documents=texts, metadatas=[{"source": s} for s in sources], embeddings=embeddings)
        return f"Index built with {len(texts)} chunks."


