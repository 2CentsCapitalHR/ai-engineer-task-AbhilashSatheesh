from typing import List
from docx import Document


def read_docx_text(path: str) -> str:
    doc = Document(path)
    parts: List[str] = []
    for p in doc.paragraphs:
        parts.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join([t for t in parts if t and t.strip()])


