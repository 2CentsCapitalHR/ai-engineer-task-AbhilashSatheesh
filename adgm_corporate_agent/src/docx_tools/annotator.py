from typing import List, Dict, Any
from docx import Document
from docx.enum.text import WD_COLOR_INDEX


def annotate_docx_with_issues(input_path: str, issues: List[Dict[str, Any]], output_path: str) -> None:
    doc = Document(input_path)

    # Build a simple index: look for mentioned clause/section text in paragraphs and highlight
    for issue in issues:
        snippet = issue.get("snippet") or issue.get("issue") or ""
        if not snippet:
            continue
        lowered = snippet[:100].lower()
        for para in doc.paragraphs:
            if lowered and lowered in para.text.lower():
                for run in para.runs:
                    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                break

    # Append review notes section
    doc.add_page_break()
    doc.add_heading("Review Notes", level=1)
    for i, issue in enumerate(issues, start=1):
        doc.add_paragraph(f"{i}. Document: {issue.get('document', 'Unknown')}")
        if issue.get("section"):
            doc.add_paragraph(f"   Section: {issue['section']}")
        doc.add_paragraph(f"   Issue: {issue.get('issue', '')}")
        if issue.get("severity"):
            doc.add_paragraph(f"   Severity: {issue['severity']}")
        if issue.get("suggestion"):
            doc.add_paragraph(f"   Suggestion: {issue['suggestion']}")
        sources = issue.get("source_citations") or []
        for s_idx, src in enumerate(sources[:3], start=1):
            doc.add_paragraph(f"   Source {s_idx}: {src.get('source')} — {src.get('snippet')}")

    doc.save(output_path)


